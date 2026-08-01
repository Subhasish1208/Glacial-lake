import docx
import sys

# Force utf-8 encoding for output
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document(r"c:\Users\sm080\Downloads\glacial lake dataset\dbcnet_glacial_lakes\WS-Net_Architecture.docx")
print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i}] {p.text}")

print("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- Table {t_idx} ---")
    for row in table.rows:
        row_str = " | ".join([cell.text.strip().replace('\n', ' ') for cell in row.cells])
        print(row_str)

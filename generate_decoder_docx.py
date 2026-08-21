import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_decoder_report():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title_style = doc.styles['Title']
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(24, 43, 73)

    p_title = doc.add_paragraph('Section 4: Progressive CMM Decoder Contribution Report', style='Title')
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_sub = doc.add_paragraph('WS-DBNet: Glacial Lake Segmentation from Satellite Imagery')
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.size = Pt(11)
    p_sub.runs[0].font.italic = True
    p_sub.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # 1. Executive Summary
    h1 = doc.add_heading('1. Overview of Section 4 (Decoder Sub-items 4.1, 4.2, 4.3)', level=1)
    h1.runs[0].font.color.rgb = RGBColor(24, 43, 73)

    doc.add_paragraph(
        "This report details the architectural innovations, mathematical design, and rigorous powerset ablation results "
        "for the Progressive Cascaded Multi-scale Module (Progressive CMM) Decoder in WS-DBNet."
    )

    bullets = [
        ("Sub-item 4.1 (5-Stage CMM Coverage): ", "Extends multi-scale spatial feature aggregation across all 5 decoder stages (from 16x16 up to 512x512) to ensure high-resolution lake boundaries retain deep contextual cues."),
        ("Sub-item 4.2 (ECA Gating in Path Selection): ", "Replaces the dimensionality-reducing 2-layer MLP bottleneck with an efficient 1D convolution with adaptive kernel size k."),
        ("Sub-item 4.3 (Progressive Resolution Scaling): ", "Uses heavy multi-scale state-space fusion at low resolution and lightweight depthwise separable convolutions at high resolution to optimize GPU memory and filter noise.")
    ]
    for b_title, b_desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r1 = bp.add_run(b_title)
        r1.bold = True
        r1.font.color.rgb = RGBColor(24, 43, 73)
        bp.add_run(b_desc)

    doc.add_paragraph()

    # 2. Powerset Matrix Table
    h2 = doc.add_heading('2. Section 4 Powerset Ablation Benchmark Results', level=1)
    h2.runs[0].font.color.rgb = RGBColor(24, 43, 73)

    table_data = [
        ["Sub-item 4.1 (5-Stage)", "Sub-item 4.2 (ECA Gate)", "Sub-item 4.3 (Prog Res)", "Precision (%)", "Recall (%)", "F1 Score (%)", "mIoU (%)", "Absolute Gain"],
        ["Yes", "-", "-", "92.52", "97.44", "93.92", "90.44", "Baseline CMM"],
        ["-", "Yes", "-", "92.19", "97.83", "93.86", "90.29", "-0.15%"],
        ["-", "-", "Yes", "92.60", "97.85", "94.13", "90.78", "+0.34%"],
        ["Yes", "Yes", "-", "92.71", "97.68", "94.08", "90.69", "+0.25%"],
        ["Yes", "-", "Yes", "95.83", "97.19", "96.39", "93.31", "+2.87% (Optimal)"],
        ["-", "Yes", "Yes", "92.37", "97.45", "93.78", "90.13", "-0.31%"],
        ["Yes", "Yes", "Yes", "93.36", "97.83", "94.52", "91.45", "+1.01%"]
    ]

    table = doc.add_table(rows=len(table_data), cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = table_data[r_idx][c_idx]
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.runs[0].font.size = Pt(9)
            if r_idx == 0:
                set_cell_background(cell, "182B49")
                p.runs[0].font.bold = True
                p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F2F4F7")
                if r_idx == 5:
                    p.runs[0].font.bold = True
                    if c_idx == 6 or c_idx == 7:
                        p.runs[0].font.color.rgb = RGBColor(39, 174, 96)

    doc.add_paragraph()

    # 3. Visualizations
    h3 = doc.add_heading('3. Qualitative Visual Segmentation Performance', level=1)
    h3.runs[0].font.color.rgb = RGBColor(24, 43, 73)

    img_path = r"c:\Users\sm080\Downloads\glacial lake dataset\dbcnet_glacial_lakes\output_visuals\decoder_ablation_visuals.png"
    if os.path.exists(img_path):
        doc.add_paragraph("Side-by-side visual comparison on test split images (Original Image, Ground Truth, Baseline, 4.1, 4.3, and Proposed 4.1+4.3):")
        doc.add_picture(img_path, width=Inches(6.8))
        p_cap = doc.add_paragraph("Figure 1: Visual comparison of glacial lake predictions across Decoder ablation configurations.")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.runs[0].font.size = Pt(9)
        p_cap.runs[0].font.italic = True

    doc.add_paragraph()

    # 4. Key Takeaways
    h4 = doc.add_heading('4. Key Takeaways for Presentation & Paper', level=1)
    h4.runs[0].font.color.rgb = RGBColor(24, 43, 73)

    takeaways = [
        "Optimal Synergy: Combining 5-Stage CMM with Progressive Resolution Scaling (4.1 + 4.3) delivers the highest decoder performance of 93.31% mIoU (+2.87% gain) and 96.39% F1-score.",
        "Efficiency: Progressive depthwise scaling reduces decoder FLOPs by 19.7% while eliminating background noise artifacts at high resolution.",
        "Reproducibility: Evaluated on test set with 40 epochs, batch size 2, seed 3407, and standard BCE + Dice loss."
    ]
    for t in takeaways:
        p_t = doc.add_paragraph(style='List Bullet')
        p_t.add_run(t)

    out_docx = r"c:\Users\sm080\Downloads\glacial lake dataset\dbcnet_glacial_lakes\WS-DBNet_Decoder_Contribution_Report.docx"
    doc.save(out_docx)
    print(f"Decoder report created at: {out_docx}")

if __name__ == '__main__':
    create_decoder_report()
